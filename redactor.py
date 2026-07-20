import re
import os
import csv

# Import optional packages
try:
    import docx
except ImportError:
    docx = None

try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None

try:
    import openpyxl
except ImportError:
    openpyxl = None

# Regex patterns for 11 PII categories
PATTERNS = {
    'email': re.compile(r'\b[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+\b'),
    'phone': re.compile(r'\b(?:\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}\b'),
    'date': re.compile(r'\b(?:\d{1,2}[-/.]\d{1,2}[-/.]\d{2,4}|\d{4}[-/.]\d{1,2}[-/.]\d{1,2})\b'),
    'aadhaar': re.compile(r'\b[2-9]\d{3}[-\s]?\d{4}[-\s]?\d{4}\b'),
    'pan': re.compile(r'\b[A-Za-z]{5}\d{4}[A-Za-z]\b'),
    'passport': re.compile(r'\b[A-Za-z]\d{7}\b'),
    'dl': re.compile(r'\b[A-Za-z]{2}[-\s]?\d{2}[-\s]?\d{4}[-\s]?\d{7}\b'),
    'bank': re.compile(r'\b\d{9,18}\b'),
    'ifsc': re.compile(r'\b[A-Za-z]{4}0[A-Za-z0-9]{6}\b'),
    'credit_card': re.compile(r'\b(?:\d{4}[-\s]?){3}\d{4}\b|\b\d{13,19}\b'),
    'upi': re.compile(r'\b[a-zA-Z0-9.\-_]+@[a-zA-Z0-9\-]+(?!\.[a-zA-Z]{2,})\b')
}

def get_masked_value(matched_text, style, custom_label=None):
    """Returns the masked version of matched_text based on selected style."""
    if not matched_text:
        return ""
    if style == 'blackout':
        return '█' * len(matched_text)
    elif style == 'asterisk':
        return '*' * len(matched_text)
    elif style == 'cross':
        return 'X' * len(matched_text)
    else:  # 'custom' or default
        return custom_label if custom_label else '[REDACTED]'

def redact_text_content(content, active_patterns, custom_terms, style='custom', 
                        custom_label='[REDACTED]', redact_all=True, case_sensitive=False):
    """
    Scans and redacts raw text content.
    Returns (redacted_content, total_count, category_counts)
    """
    category_counts = {cat: 0 for cat in PATTERNS}
    category_counts['manual'] = 0
    total_count = 0
    sub_limit = 0 if redact_all else 1

    # Redact predefined patterns
    for cat_name, regex in PATTERNS.items():
        if active_patterns.get(cat_name, False):
            # Find matches first to count them
            matches = regex.findall(content)
            # Filter phone matches to keep only the longest digit sequences
            if cat_name == 'phone':
                max_len = 0
                digit_lengths = []
                for m in matches:
                    digits = re.sub(r'\\D', '', m)
                    dlen = len(digits)
                    digit_lengths.append((m, dlen))
                    if dlen > max_len:
                        max_len = dlen
                matches = [m for m, dlen in digit_lengths if dlen == max_len]
            unique_matches = set(matches) if redact_all else matches[:1]

            for match in unique_matches:
                if isinstance(match, tuple):
                    match = match[0]
                match = match.strip()
                if not match:
                    continue

                # Perform replace
                pattern = re.compile(re.escape(match))
                masked = get_masked_value(match, style, custom_label)
                content, num_subs = pattern.subn(masked, content, count=sub_limit)

                category_counts[cat_name] += num_subs
                total_count += num_subs

    # Redact manual custom terms
    for term in custom_terms:
        term = term.strip()
        if not term:
            continue
        
        flags = 0 if case_sensitive else re.IGNORECASE
        pattern = re.compile(re.escape(term), flags)
        
        masked = get_masked_value(term, style, custom_label)
        content, num_subs = pattern.subn(masked, content, count=sub_limit)
        
        category_counts['manual'] += num_subs
        total_count += num_subs
        
    return content, total_count, category_counts

def redact_file(input_path, output_path, active_patterns, custom_terms, 
                style='custom', custom_label='[REDACTED]', redact_all=True, 
                case_sensitive=False, cells=None, rows=None, cols=None):
    """
    Parses, redacts, and saves a file based on format.
    Supports TXT, CSV, JSON, DOCX, Excel, and PDF.
    Returns (success, total_count, category_counts, error_message)
    """
    _, ext = os.path.splitext(input_path.lower())
    category_counts = {cat: 0 for cat in PATTERNS}
    category_counts['manual'] = 0
    total_count = 0
    sub_limit = 0 if redact_all else 1
    
    cells = cells or []
    rows = rows or []  # 0-indexed integers
    cols = cols or []  # 0-indexed integers (col index)
    
    try:
        # --- 1. Plain Text / JSON ---
        if ext in ['.txt', '.json', '.xml']:
            with open(input_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
            
            redacted_content, count, counts = redact_text_content(
                content, active_patterns, custom_terms, style, custom_label, redact_all, case_sensitive
            )
            
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(redacted_content)
                
            return True, count, counts, None

        # --- 2. CSV / Spreadsheet Grid ---
        elif ext == '.csv':
            # Read all rows
            with open(input_path, 'r', encoding='utf-8', errors='ignore') as f:
                reader = list(csv.reader(f))
                
            count = 0
            # Redact cells based on coordinates (e.g. cells like 'A1', row index, col index)
            for r_idx, row in enumerate(reader):
                for c_idx, val in enumerate(row):
                    # Check if cell/row/col selected
                    col_letter = chr(65 + c_idx) if c_idx < 26 else f"Col{c_idx}"
                    cell_coord = f"{col_letter}{r_idx + 1}"
                    
                    if (cell_coord in cells) or (r_idx in rows) or (c_idx in cols):
                        reader[r_idx][c_idx] = get_masked_value(val, style, custom_label)
                        category_counts['manual'] += 1
                        count += 1
                    else:
                        # Otherwise run standard text redactions on the cell content
                        masked_val, cell_count, cell_counts = redact_text_content(
                            val, active_patterns, custom_terms, style, custom_label, redact_all, case_sensitive
                        )
                        reader[r_idx][c_idx] = masked_val
                        count += cell_count
                        for cat in category_counts:
                            category_counts[cat] += cell_counts[cat]
            
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            with open(output_path, 'w', newline='', encoding='utf-8') as f:
                writer = csv.writer(f)
                writer.writerows(reader)
                
            return True, count, category_counts, None

        # --- 3. Excel Documents (XLSX) ---
        elif ext == '.xlsx':
            if not openpyxl:
                return False, 0, category_counts, "openpyxl library not available"
                
            wb = openpyxl.load_workbook(input_path)
            sheet = wb.active
            count = 0
            
            for row in sheet.iter_rows():
                for cell in row:
                    if cell.value is not None:
                        val_str = str(cell.value)
                        r_idx = cell.row - 1      # 0-indexed
                        c_idx = cell.column - 1   # 0-indexed
                        cell_coord = cell.coordinate
                        
                        if (cell_coord in cells) or (r_idx in rows) or (c_idx in cols):
                            cell.value = get_masked_value(val_str, style, custom_label)
                            category_counts['manual'] += 1
                            count += 1
                        else:
                            masked_val, cell_count, cell_counts = redact_text_content(
                                val_str, active_patterns, custom_terms, style, custom_label, redact_all, case_sensitive
                            )
                            cell.value = masked_val
                            count += cell_count
                            for cat in category_counts:
                                category_counts[cat] += cell_counts[cat]
                                
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            wb.save(output_path)
            return True, count, category_counts, None

        # --- 4. Word Documents (DOCX) ---
        elif ext == '.docx':
            if not docx:
                return False, 0, category_counts, "docx library not available"
                
            doc = docx.Document(input_path)
            count = 0
            
            def redact_runs(runs):
                nonlocal count
                for run in runs:
                    if run.text:
                        red_text, run_count, run_counts = redact_text_content(
                            run.text, active_patterns, custom_terms, style, custom_label, redact_all, case_sensitive
                        )
                        run.text = red_text
                        count += run_count
                        for cat in category_counts:
                            category_counts[cat] += run_counts[cat]
                            
            # Redact paragraphs
            for paragraph in doc.paragraphs:
                redact_runs(paragraph.runs)
                
            # Redact tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            redact_runs(paragraph.runs)
                            
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            doc.save(output_path)
            return True, count, category_counts, None

        # --- 5. PDF Documents ---
        elif ext == '.pdf':
            if not fitz:
                return False, 0, category_counts, "pymupdf library not available"
                
            doc = fitz.open(input_path)
            count = 0
            
            # Formulate text terms to search
            # We compile active patterns
            active_regexes = []
            for name, regex in PATTERNS.items():
                if active_patterns.get(name, False):
                    active_regexes.append((name, regex))
            
            for page in doc:
                page_text = page.get_text()
                terms_to_redact = []
                
                # Predefined regexes lookup
                for name, regex in active_regexes:
                    matches = regex.findall(page_text)
                    for m in set(matches):
                        if isinstance(m, tuple):
                            m = m[0]
                        m = m.strip()
                        if m:
                            terms_to_redact.append((m, name))
                            
                # Custom keywords lookup
                for term in custom_terms:
                    term = term.strip()
                    if term:
                        flags = re.IGNORECASE if not case_sensitive else 0
                        matches = re.findall(re.escape(term), page_text, flags)
                        for m in set(matches):
                            if m.strip():
                                terms_to_redact.append((m, 'manual'))
                
                # If Only first occurrence flag is active, we slice terms
                if not redact_all and terms_to_redact:
                    terms_to_redact = terms_to_redact[:1]
                    
                # Redact coordinates
                for term, cat in terms_to_redact:
                    rects = page.search_for(term)
                    for rect in rects:
                        # In PDF, we black-box if blackout style, or draw mask texts
                        # Fitz's add_redact_annot allows custom replacement text!
                        # We specify the replacement text corresponding to the mask style!
                        mask_text = get_masked_value(term, style, custom_label)
                        page.add_redact_annot(rect, text=mask_text, fill=(0, 0, 0) if style == 'blackout' else None)
                        category_counts[cat] += 1
                        count += 1
                        
                page.apply_redactions()
                
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            doc.save(output_path)
            doc.close()
            return True, count, category_counts, None
            
        else:
            return False, 0, category_counts, f"Unsupported file type: {ext}"
            
    except Exception as e:
        return False, 0, category_counts, str(e)
